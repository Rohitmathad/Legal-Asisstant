
from flask import Flask, request, jsonify, send_file
import requests
import MySQLdb
from flask_cors import CORS
from docx import Document
import mammoth
import psycopg2
from dotenv import load_dotenv
import os
import sys

current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)

# Importing both models
from model.bot import get_response
from model.similarity import get_document

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

# Configure CORS to allow requests from React frontend
cors = CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})

db = psycopg2.connect(database=os.getenv('DATABASE_NAME'), user=os.getenv('DATABASE_USER'),
                      password=os.getenv('PASSWORD'), host=os.getenv('DATABASE_HOST'), port=os.getenv('DATABASE_PORT'), keepalives=1, keepalives_idle=30,
                      keepalives_interval=10, keepalives_count=5)

# Get all the services
@app.route('/api/services', methods=["GET"])
def services():
    cur = db.cursor()
    cur.execute('SELECT * FROM services')
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    print(json_data)
    return jsonify(json_data)

# Get forms of a particular service
@app.route('/api/forms', methods=["GET"])
def get_forms():
    # Send json object {"service_id": "..."}
    Service = request.args.get('service_id')
    print(type(Service))
    print(Service)
    cur = db.cursor()

    # Update the SQL query to only include service IDs up to 3
    cur.execute(
        "SELECT services.service_id, services.service_name, forms.form_id, forms.form_name, forms.form_link FROM services INNER JOIN forms ON services.service_id = forms.service_id WHERE forms.service_id = %s AND services.service_id <= 3;", [Service]
    )

    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    print(json_data)
    return jsonify(json_data)

# Get all queries for a form
@app.route('/api/form-details', methods=["GET"])
def get_form_details():
    # Send json object {"form_id":"..."}
    form_id = request.args.get('form_id')
    print(form_id)
    cur = db.cursor()
    cur.execute("SELECT * FROM forms WHERE form_id = %s;", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))

    cur.execute("SELECT * FROM ques_categories WHERE id IN (SELECT DISTINCT(category_id) FROM input_ques WHERE ques_id IN (SELECT form_query_id FROM form_queries WHERE form_id = %s));", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.execute("SELECT * FROM input_ques WHERE ques_id IN (SELECT form_query_id FROM form_queries WHERE form_id = %s);", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    return jsonify(json_data)

# Return the contents of final doc
@app.route('/api/final-content', methods=["POST"])
def final_content():
    form_details = request.json  # Under Progress
    form_id = form_details["form_id"]
    # print(type(form_details))
    print(form_id)
    cur = db.cursor()
    cur.execute("SELECT form_link FROM forms where form_id = %s;", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    print(json_data[0]["form_link"])
    response = requests.get(json_data[0]["form_link"])
    directory = './docs'

    if not os.path.exists(directory):
        os.makedirs(directory)

    file_path = './docs/localfile.docx'

    with open(file_path, 'wb') as f:
        f.write(response.content)

    doc = Document('./docs/localfile.docx')
    test = list([int(x) for x in form_details.keys() if x.isdigit()])

    test.sort(reverse=True)
    print(test)
    for key in test:
        old = '#'+str(key)
        new = str(form_details[str(key)])

        for p in doc.paragraphs:
            if old in p.text:
                # print(old)
                inline = p.runs
                for i in range(len(inline)):
                    if old in inline[i].text:
                        # print(old)
                        res = inline[i].text.replace(old, new)
                        inline[i].text = res
    doc.save("./docs/Output2.docx")

    f = open('./docs/Output2.docx', 'rb')

    docx_content = mammoth.convert_to_html(f)
    # print(docx_content.value)
    # docx_content.close()

    # fullText = []
    # for para in doc.paragraphs:
    #     fullText.append(para.text)
    # fullText = '\n'.join(fullText)
    # print(fullText)
    return jsonify({'content': docx_content.value})

# Return the final doc
@app.route('/api/final-form', methods=["POST"])
def final_form():
    contents = request.get_json()
    print(contents)
    with open('docs/Output2.docx', 'w') as file:
        file.write(contents)
    # doc = Document('docs/localfile.docx')
    # for key, value in form_details.items():
    #     for paragraph in doc.paragraphs:
    #         paragraph.text = paragraph.text.replace(
    #             "#"+str(key)+'#', str(value))

    # doc.save("docs/Output2.docx")
    return send_file('./docs/Output2.docx', as_attachment=True)

@app.route('/api/chat', methods=['POST'])
def chat():
    user_input = request.json

    # Choose one of the following models. get_response works using Bag of Words Principle while get_document works using Cosine Similarity
    # response = get_response(user_input['user_chat'])
    response = get_document(user_input['user_chat'])
    return jsonify({'aiMessage': response})

if __name__ == '__main__':
    app.run(debug=True)




# from os import environ, path
# from typing import List
# import chromadb
# from dotenv import load_dotenv
# from langchain.chains.combine_documents.stuff import create_stuff_documents_chain
# from langchain.chains.retrieval import create_retrieval_chain
# from langchain.docstore.document import Document
# from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader
# from langchain_community.embeddings import JinaEmbeddings
# from langchain_community.vectorstores.chroma import Chroma
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.runnables import Runnable
# from langchain_core.vectorstores import VectorStoreRetriever
# from langchain_groq import ChatGroq
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from termcolor import cprint
# from transformers import AutoTokenizer
# import mysql.connector

# # Flask imports
# from flask import Flask, request, jsonify, send_file
# import requests
# import MySQLdb
# from flask_cors import CORS
# from docx import Document
# import mammoth
# import psycopg2
# from dotenv import load_dotenv
# import os
# import sys


# current_dir = os.path.dirname(os.path.abspath(__file__))
# parent_dir = os.path.dirname(current_dir)
# sys.path.append(parent_dir)

# # Importing both models
# from model.bot import get_response
# from model.similarity import get_document

# # Load environment variables from .env file
# load_dotenv()

# app = Flask(__name__)

# # Configure CORS to allow requests from React frontend
# cors = CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})

# db = psycopg2.connect(database=os.getenv('DATABASE_NAME'), user=os.getenv('DATABASE_USER'),
#                       password=os.getenv('PASSWORD'), host=os.getenv('DATABASE_HOST'), port=os.getenv('DATABASE_PORT'), keepalives=1, keepalives_idle=30,
#                       keepalives_interval=10, keepalives_count=5)

# # Get all the services
# @app.route('/api/services', methods=["GET"])
# def services():
#     cur = db.cursor()
#     cur.execute('SELECT * FROM services')
#     row_headers = [x[0] for x in cur.description]
#     rv = cur.fetchall()
#     json_data = []
#     for result in rv:
#         json_data.append(dict(zip(row_headers, result)))
#     cur.close()
#     print(json_data)
#     return jsonify(json_data)

# # Get forms of a particular service
# @app.route('/api/forms', methods=["GET"])
# def get_forms():
#     # Send json object {"service_id": "..."}
#     Service = request.args.get('service_id')
#     print(type(Service))
#     print(Service)
#     cur = db.cursor()

#     # Update the SQL query to only include service IDs up to 3
#     cur.execute(
#         "SELECT services.service_id, services.service_name, forms.form_id, forms.form_name, forms.form_link FROM services INNER JOIN forms ON services.service_id = forms.service_id WHERE forms.service_id = %s AND services.service_id <= 3;", [Service]
#     )

#     row_headers = [x[0] for x in cur.description]
#     rv = cur.fetchall()
#     json_data = []
#     for result in rv:
#         json_data.append(dict(zip(row_headers, result)))
#     cur.close()
#     print(json_data)
#     return jsonify(json_data)

# # Get all queries for a form
# @app.route('/api/form-details', methods=["GET"])
# def get_form_details():
#     # Send json object {"form_id":"..."}
#     form_id = request.args.get('form_id')
#     print(form_id)
#     cur = db.cursor()
#     cur.execute("SELECT * FROM forms WHERE form_id = %s;", [form_id])
#     row_headers = [x[0] for x in cur.description]
#     rv = cur.fetchall()
#     json_data = []
#     for result in rv:
#         json_data.append(dict(zip(row_headers, result)))

#     cur.execute("SELECT * FROM ques_categories WHERE id IN (SELECT DISTINCT(category_id) FROM input_ques WHERE ques_id IN (SELECT form_query_id FROM form_queries WHERE form_id = %s));", [form_id])
#     row_headers = [x[0] for x in cur.description]
#     rv = cur.fetchall()
#     for result in rv:
#         json_data.append(dict(zip(row_headers, result)))
#     cur.execute("SELECT * FROM input_ques WHERE ques_id IN (SELECT form_query_id FROM form_queries WHERE form_id = %s);", [form_id])
#     row_headers = [x[0] for x in cur.description]
#     rv = cur.fetchall()
#     for result in rv:
#         json_data.append(dict(zip(row_headers, result)))
#     cur.close()
#     return jsonify(json_data)

# # Return the contents of final doc
# @app.route('/api/final-content', methods=["POST"])
# def final_content():
#     form_details = request.json  # Under Progress
#     form_id = form_details["form_id"]
#     # print(type(form_details))
#     print(form_id)
#     cur = db.cursor()
#     cur.execute("SELECT form_link FROM forms where form_id = %s;", [form_id])
#     row_headers = [x[0] for x in cur.description]
#     rv = cur.fetchall()
#     json_data = []
#     for result in rv:
#         json_data.append(dict(zip(row_headers, result)))
#     cur.close()
#     print(json_data[0]["form_link"])
#     response = requests.get(json_data[0]["form_link"])
#     directory = './docs'

#     if not os.path.exists(directory):
#         os.makedirs(directory)

#     file_path = './docs/localfile.docx'

#     with open(file_path, 'wb') as f:
#         f.write(response.content)

#     doc = Document('./docs/localfile.docx')
#     test = list([int(x) for x in form_details.keys() if x.isdigit()])

#     test.sort(reverse=True)
#     print(test)
#     for key in test:
#         old = '#'+str(key)
#         new = str(form_details[str(key)])

#         for p in doc.paragraphs:
#             if old in p.text:
#                 # print(old)
#                 inline = p.runs
#                 for i in range(len(inline)):
#                     if old in inline[i].text:
#                         # print(old)
#                         res = inline[i].text.replace(old, new)
#                         inline[i].text = res
#     doc.save("./docs/Output2.docx")

#     f = open('./docs/Output2.docx', 'rb')

#     docx_content = mammoth.convert_to_html(f)
#     # print(docx_content.value)
#     # docx_content.close()

#     # fullText = []
#     # for para in doc.paragraphs:
#     #     fullText.append(para.text)
#     # fullText = '\n'.join(fullText)
#     # print(fullText)
#     return jsonify({'content': docx_content.value})

# # Return the final doc
# @app.route('/api/final-form', methods=["POST"])
# def final_form():
#     contents = request.get_json()
#     print(contents)
#     with open('docs/Output2.docx', 'w') as file:
#         file.write(contents)
#     # doc = Document('docs/localfile.docx')
#     # for key, value in form_details.items():
#     #     for paragraph in doc.paragraphs:
#     #         paragraph.text = paragraph.text.replace(
#     #             "#"+str(key)+'#', str(value))

#     # doc.save("docs/Output2.docx")
#     return send_file('./docs/Output2.docx', as_attachment=True)


# # ... (The rest of the code from the first file)

# # Flask app and routes
# app = Flask(__name__)
# cors = CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})

# EMBED_MODEL_NAME = "jina-embeddings-v2-base-en"
# LLM_NAME = "mixtral-8x7b-32768"
# LLM_TEMPERATURE = 0.1

# # this is the maximum chunk size allowed by the chosen embedding model. You can choose a smaller size.
# CHUNK_SIZE = 8192

# DOCUMENT_DIR = "./documents/"  # the directory where the pdf files should be placed
# VECTOR_STORE_DIR = "./vectorstore/"  # the directory where the vectors are stored
# COLLECTION_NAME = "collection1"  # chromadb collection name
# # ===============================================================

# load_dotenv()


# def load_documents() -> List[Document]:
#     """Loads the pdf files within the DOCUMENT_DIR constant."""
#     try:
#         print("[+] Loading documents...")

#         documents = DirectoryLoader(
#             path.join(DOCUMENT_DIR), glob="**/*.pdf", loader_cls=PyPDFLoader
#         ).load()
#         cprint(f"[+] Document loaded, total pages: {len(documents)}", "green")

#         return documents
#     except:
#         cprint("[-] Error loading the document.", "red")


# def chunk_document(documents: List[Document]) -> List[Document]:
#     """Splits the input documents into maximum of CHUNK_SIZE chunks."""
#     tokenizer = AutoTokenizer.from_pretrained(
#         "jinaai/" + EMBED_MODEL_NAME, cache_dir=environ.get("HF_HOME")
#     )
#     text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
#         tokenizer=tokenizer,
#         chunk_size=CHUNK_SIZE,
#         chunk_overlap=CHUNK_SIZE // 50,
#     )

#     print(f"[+] Splitting documents...")
#     chunks = text_splitter.split_documents(documents)
#     cprint(f"[+] Document splitting done, {len(chunks)} chunks total.", "green")

#     return chunks


# def create_and_store_embeddings(
#     embedding_model: JinaEmbeddings, chunks: List[Document]
# ) -> Chroma:
#     """Calculates the embeddings and stores them in a a chroma vectorstore."""
#     vectorstore = Chroma.from_documents(
#         chunks,
#         embedding=embedding_model,
#         collection_name=COLLECTION_NAME,
#         persist_directory=VECTOR_STORE_DIR,
#     )
#     cprint("[+] Vectorstore created.", "green")

#     return vectorstore


# def get_vectorstore_retriever(embedding_model: JinaEmbeddings) -> VectorStoreRetriever:
#     """Returns the vectorstore."""
#     db = chromadb.PersistentClient(VECTOR_STORE_DIR)
#     try:
#         # Check for the existence of the vectorstore specified by the COLLECTION_NAME
#         db.get_collection(COLLECTION_NAME)
#         retriever = Chroma(
#             embedding_function=embedding_model,
#             collection_name=COLLECTION_NAME,
#             persist_directory=VECTOR_STORE_DIR,
#         ).as_retriever(search_kwargs={"k": 3})
#     except:
#         # The vectorstore doesn't exist, so create it.
#         pdf = load_documents()
#         chunks = chunk_document(pdf)
#         retriever = create_and_store_embeddings(embedding_model, chunks).as_retriever(
#             search_kwargs={"k": 3}
#         )

#     return retriever


# def create_rag_chain(embedding_model: JinaEmbeddings, llm: ChatGroq) -> Runnable:
#     """Creates the RAG chain."""
#     template = """ Act as an Indian legal assistant , answer queries related only to that and deny answering other irrelevant questions 
#     {context}
#     </context>

#     Question: {input}
#     """
#     prompt = ChatPromptTemplate.from_template(template)

#     document_chain = create_stuff_documents_chain(llm=llm, prompt=prompt)

#     retriever = get_vectorstore_retriever(embedding_model)

#     retrieval_chain = create_retrieval_chain(retriever, document_chain)

#     return retrieval_chain


# def run_chain(chain: Runnable) -> None:
#     """Run the RAG chain with the user query."""
#     while True:
#         query = input("Enter a prompt: ")
#         if query.lower() in ["q", "quit", "exit"]:
#             return
#         response = chain.invoke({"input": query})

#         for doc in response["context"]:
#             cprint(
#                 f"[+] {doc.metadata} | content: {doc.page_content[:20]}...",
#                 "light_yellow",
#             )

#         cprint("\n" + response["answer"], end="\n\n", color="light_blue")


# def main() -> None:
#     embedding_model = JinaEmbeddings(
#         jina_api_key=environ.get("JINA_API_KEY"),
#         model_name=EMBED_MODEL_NAME,
#     )

#     llm = ChatGroq(temperature=LLM_TEMPERATURE, model_name=LLM_NAME)

#     chain = create_rag_chain(embedding_model=embedding_model, llm=llm)

#     run_chain(chain)

# # ... (The rest of the code from the second file)

# # Chatbot API route
# @app.route('/api/chat', methods=['POST'])
# def chat():
#     user_input = request.json

#     # Choose one of the following models. get_response works using Bag of Words Principle while get_document works using Cosine Similarity
#     # response = get_response(user_input['user_chat'])
#     response = get_document(user_input['user_chat'])
#     return jsonify({'aiMessage': response})

# if __name__ == '__main__':
#     app.run(debug=True)
